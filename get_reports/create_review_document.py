"""Module to create a human readable report for expedited review."""

import argparse
import json
import logging
import os
from docx import Document


def _parse_arguments() -> argparse.Namespace:
    """Parse the command line arguments."""
    arg_parser = argparse.ArgumentParser(
        prog="create_review_document",
        description="Create a human readable report for expedited review.",
    )
    arg_parser.add_argument(
        "--input",
        help="Observations requiring review",
        default="reports/records_to_review.json",
    )

    arg_parser.add_argument(
        "--output",
        help="Observations requiring review",
        default="reports/records_to_review.docx",
    )

    arg_parser.add_argument(
        "--verbose", action="store_true", help="increase verbosity"
    )
    return arg_parser.parse_args()


def _load_observations(file_path: str) -> dict:
    """Load observations from a JSON file."""
    with open(file_path, "rt", encoding="utf-8") as file:
        return json.load(file)


def _create_document(observations: dict) -> Document:
    """Create a Word document based on observations."""
    document = Document()
    _add_document_header(document, observations)
    _add_county_records(document, observations["records"])
    return document


LIST_BULLET_STYLE = "List Bullet"


def _add_document_header(document: Document, observations: dict):
    """Add the header section to the document."""
    document.add_heading("DRAFT Records for Expedited Review", 0)
    document.add_paragraph(
        style=LIST_BULLET_STYLE,
        text=f"Dates of observations: {observations['date of observations']}",
    )
    document.add_paragraph(
        style=LIST_BULLET_STYLE,
        text=f"Region: {observations['region']}",
    )
    document.add_paragraph(
        style=LIST_BULLET_STYLE,
        text=f"Dates of report creation: {observations['date of report']}",
    )
    p = document.add_paragraph(
        style=LIST_BULLET_STYLE,
        text="Produced for and by VARCOM (Virginia Avian Records Committee)"
        ", of the VSO (Virginia Society of Ornithology). ",
    )
    p.add_run(
        "https://www.virginiabirds.org/varcom"
    ).hyperlink = "https://www.virginiabirds.org/varcom"
    p = document.add_paragraph(
        style=LIST_BULLET_STYLE,
        text="Automatically generated by VARCOM-get-reports. ",
    )
    p.add_run(
        "https://github.com/gbabineau/VARCOM-get-reports"
    ).hyperlink = "https://github.com/gbabineau/VARCOM-get-reports"


def _add_county_records(document: Document, counties: list):
    """Add records for each county to the document."""
    for county in counties:
        document.add_heading(
            f"Reviewable records in {county['county']}", level=1
        )
        document.add_paragraph(f"Total records: {len(county['records'])}")
        _add_species_records(document, county)


def _add_species_records(document: Document, county: dict):
    """Add species records for a specific county."""
    county_records = county["records"]
    sorted_records = sorted(
        county_records,
        key=lambda x: (
            x["observation"]["comName"],
            x["observation"]["obsDt"],
        ),
    )

    current_species = ""
    for record in sorted_records:
        species = record["observation"]["comName"]
        if species != current_species:
            _add_species_heading(document, species, record, county)
            current_species = species
        if record["observation"].get("subId"):
            _add_observation_data(document, record)


def _add_species_heading(
    document: Document, species: str, record: dict, county: dict
):
    """Add a heading and details for a species."""
    document.add_heading(species, level=2)
    review_species = record.get("review_species", {})
    if exclude := review_species.get("exclude", []):
        document.add_paragraph(
            f"The species {species} is not excluded from review in {county['county']} because it is not in the following counties or groups of counties: {exclude}"
        )
    if only := review_species.get("only", []):
        document.add_paragraph(
            f"The species {species}: is only reviewed in the following counties or groups of counties: {only}"
        )
    if unique_exclude_notes := review_species.get("uniqueExcludeNotes", None):
        document.add_paragraph(
            f"This species has unique Exclude Notes which could not be automated. {unique_exclude_notes}"
        )
    if not exclude and not only and not unique_exclude_notes:
        document.add_paragraph(
            f"The species {species} is reviewable across the entire state."
        )
    if record.get("new", False):
        document.add_paragraph(
            f"The species {species} is not in the state list. A new record?"
        )


def _add_observation_data(document: Document, record: dict):
    """Add a hyperlink for a species record."""
    observation = record['observation']
    p = document.add_paragraph(style="List Bullet")
    media_status = "Has media." if record.get('media') else "No media."
    p.add_run(
        f"{observation['obsDt']}, {media_status}, Checklist: https://ebird.org/checklist/{observation['subId']}"
    ).hyperlink = f"https://ebird.org/checklist/{observation['subId']}"


def _save_document(document: Document, output: str):
    """Save the document to a file."""
    if os.path.exists(output):
        os.remove(output)
    document.save(output)


def main():
    """Main function for the app."""
    args = _parse_arguments()

    if args.verbose:
        logging.basicConfig(level=logging.INFO)

    observations = _load_observations(args.input)
    document = _create_document(observations)
    _save_document(document, args.output)


if __name__ == "__main__":
    main()
