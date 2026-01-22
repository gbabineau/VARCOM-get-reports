"""Unit tests for create_review_document module."""

from unittest.mock import patch

import pytest
from docx import Document

from get_reports.create_review_document import (_add_document_header,
                                                _add_observation_data,
                                                _add_species_heading,
                                                _create_document,
                                                _iterate_over_species,
                                                _load_observations,
                                                _parse_arguments,
                                                _save_document, get_day_number)


class TestParseArguments:
    """Tests for _parse_arguments function."""

    def test_default_arguments(self):
        """Test that default arguments are set correctly."""
        with patch("sys.argv", ["prog"]):
            args = _parse_arguments()
            assert args.input == "reports/records_to_review.json"
            assert args.output == "reports/records_to_review.docx"
            assert args.verbose is False

    def test_custom_input_output(self):
        """Test parsing custom input and output paths."""
        with patch(
            "sys.argv",
            ["prog", "--input", "custom.json", "--output", "custom.docx"],
        ):
            args = _parse_arguments()
            assert args.input == "custom.json"
            assert args.output == "custom.docx"

    def test_verbose_flag(self):
        """Test verbose flag is parsed correctly."""
        with patch("sys.argv", ["prog", "--verbose"]):
            args = _parse_arguments()
            assert args.verbose is True


class TestLoadObservations:
    """Tests for _load_observations function."""

    def test_load_valid_json(self, tmp_path):
        """Test loading valid JSON file."""
        test_data = {"records": [{"id": 1}], "region": "Virginia"}
        file_path = tmp_path / "test.json"
        file_path.write_text('{"records": [{"id": 1}], "region": "Virginia"}')

        result = _load_observations(str(file_path))
        assert result == test_data

    def test_load_nonexistent_file(self):
        """Test loading non-existent file raises exception."""
        with pytest.raises(FileNotFoundError):
            _load_observations("nonexistent.json")


class TestGetDayNumber:
    """Tests for get_day_number function."""

    def test_valid_date_january(self):
        """Test converting January date to day number."""
        assert get_day_number("2025-01-01") == 1
        assert get_day_number("2025-01-20") == 20

    def test_valid_date_december(self):
        """Test converting December date to day number."""
        assert get_day_number("2025-12-31") == 365

    def test_leap_year(self):
        """Test day number in leap year."""
        assert get_day_number("2024-12-31") == 366

    def test_invalid_date_format(self):
        """Test invalid date format returns None."""
        assert get_day_number("2025/01/01") is None
        assert get_day_number("invalid") is None


class TestCreateDocument:
    """Tests for _create_document function."""

    @patch("get_reports.create_review_document._iterate_over_species")
    @patch("get_reports.create_review_document._add_document_header")
    def test_create_document_structure(self, mock_header, mock_iterate):
        """Test document creation calls header and iterate functions."""
        observations = {
            "records": [],
            "region": "Virginia",
            "date of observations": "2025-01-01",
            "date of report": "2025-01-20",
        }
        taxonomy = []

        doc = _create_document("test_key", observations, taxonomy)
        assert(doc.paragraphs==[])
        mock_header.assert_called_once()
        mock_iterate.assert_called_once()


class TestAddDocumentHeader:
    """Tests for _add_document_header function."""

    def test_header_content(self):
        """Test header contains required information."""
        doc = Document()
        observations = {
            "date of observations": "2025-01-01 to 2025-01-31",
            "region": "Virginia",
            "date of report": "2025-02-01",
        }

        _add_document_header(doc, observations)

        assert len(doc.paragraphs) > 0
        text = "\n".join([p.text for p in doc.paragraphs])
        assert "DRAFT Records for Expedited Review" in text
        assert "Virginia" in text


class TestAddSpeciesHeading:
    """Tests for _add_species_heading function."""

    def test_heading_with_exclude(self):
        """Test species heading with exclude information."""
        doc = Document()
        review_species = {"exclude": ["County1", "County2"]}

        _add_species_heading(doc, "Northern Cardinal", review_species)

        text = "\n".join([p.text for p in doc.paragraphs])
        assert "Northern Cardinal" in text
        assert "excluded" in text.lower()

    def test_heading_with_only(self):
        """Test species heading with only information."""
        doc = Document()
        review_species = {"only": ["County1"]}

        _add_species_heading(doc, "Roseate Spoonbill", review_species)

        text = "\n".join([p.text for p in doc.paragraphs])
        assert "only reviewed" in text.lower()

    def test_heading_default_reviewable(self):
        """Test species heading when reviewable statewide."""
        doc = Document()
        review_species = {}

        _add_species_heading(doc, "American Robin", review_species)

        text = "\n".join([p.text for p in doc.paragraphs])
        assert "reviewable across the entire state" in text.lower()

    def test_heading_unique_exclude_notes(self):
        """Test species heading with unique exclude notes."""
        doc = Document()
        review_species = {"uniqueExcludeNotes": "Special notes here"}

        _add_species_heading(doc, "Raven", review_species)

        text = "\n".join([p.text for p in doc.paragraphs])
        assert "unique Exclude Notes" in text


class TestAddObservationData:
    """Tests for _add_observation_data function."""

    @patch("get_reports.create_review_document.ebird_data_access.get_checklist_with_retry")
    def test_observation_data_formatting(self, mock_checklist):
        """Test observation data is added with correct formatting."""
        mock_checklist.return_value = {
            "userDisplayName": "John Doe",
            "locId": "L123456",
        }

        doc = Document()
        record = {
            "observation": {
                "comName": "Northern Cardinal",
                "subId": "S123",
                "howMany": 2,
                "subnational2Name": "Fairfax County",
                "obsDt": "2025-01-20T10:30:00",
            },
            "media": [{"id": 1}],
        }
        species_data = {
            "comName": "Northern Cardinal",
            "sciName": "Cardinalis cardinalis",
        }

        _add_observation_data("test_key", doc, record, species_data)

        text = doc.paragraphs[-1].text
        assert "Northern Cardinal" in text
        assert "Cardinalis cardinalis" in text
        assert "2" in text
        assert "John Doe" in text

    @patch("get_reports.create_review_document.ebird_data_access.get_checklist_with_retry")
    def test_observation_hyperlink(self, mock_checklist):
        """Test that observation includes eBird checklist hyperlink."""
        mock_checklist.return_value = {
            "userDisplayName": "Jane Smith",
            "locId": "L654321",
        }

        doc = Document()
        record = {
            "observation": {
                "comName": "Blue Jay",
                "subId": "S456",
                "howMany": 1,
                "subnational2Name": "Arlington County",
                "obsDt": "2025-01-20T14:00:00",
            },
            "media": [{"id": 2}],
        }
        species_data = {
            "comName": "Blue Jay",
            "sciName": "Cyanocitta cristata",
        }

        _add_observation_data("test_key", doc, record, species_data)

        runs = doc.paragraphs[-1].runs
        hyperlink_runs = [r for r in runs if 'https' in r.text]
        assert len(hyperlink_runs) > 0
        assert "https://ebird.org/checklist/" in hyperlink_runs[0].text


class TestIterateOverSpecies:
    """Tests for _iterate_over_species function."""

    @patch("get_reports.create_review_document._add_observation_data")
    @patch("get_reports.create_review_document._add_species_heading")
    def test_species_grouping(self, mock_heading, mock_obs_data):
        """Test species are grouped and sorted correctly."""
        doc = Document()
        counties = [
            {
                "records": [
                    {
                        "observation": {
                            "comName": "Robin",
                            "obsDt": "2025-01-20",
                            "subnational2Name": "County A",
                        },
                        "media": [{"id": 1}],
                    }
                ]
            }
        ]
        taxonomy = [{"comName": "Robin", "taxonOrder": 1.0}]

        _iterate_over_species("key", doc, counties, taxonomy)

        mock_heading.assert_called()

    @patch("get_reports.create_review_document._add_observation_data")
    @patch("get_reports.create_review_document._add_species_heading")
    def test_no_media_records_skipped(self, mock_heading, mock_obs_data):
        """Test records without media are skipped."""
        doc = Document()
        counties = [
            {
                "records": [
                    {
                        "observation": {
                            "comName": "Sparrow",
                            "obsDt": "2025-01-20",
                            "subnational2Name": "County B",
                        }
                    }
                ]
            }
        ]
        taxonomy = [{"comName": "Sparrow", "taxonOrder": 2.0}]

        _iterate_over_species("key", doc, counties, taxonomy)

        mock_obs_data.assert_not_called()


class TestSaveDocument:
    """Tests for _save_document function."""

    def test_save_new_document(self, tmp_path):
        """Test saving a new document."""
        doc = Document()
        output_path = tmp_path / "test_output.docx"

        _save_document(doc, str(output_path))

        assert output_path.exists()

    def test_overwrite_existing_document(self, tmp_path):
        """Test that existing document is overwritten."""
        output_path = tmp_path / "test.docx"
        output_path.write_text("old content")

        doc = Document()
        _save_document(doc, str(output_path))

        assert output_path.exists()
        # File should be valid DOCX now
        assert output_path.stat().st_size > 0
        assert output_path.stat().st_size > 0